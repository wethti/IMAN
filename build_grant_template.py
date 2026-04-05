"""
Grant Proposal Template — faithful replica of the Crown-IMAN 2025 document structure.

Exact specs extracted from the original docx:
  Body font:    Garamond, 12pt (sz=24), color #0D0D0D, line-spacing 259/auto
  Heading font: Aptos Display (majorHAnsi theme)
  Heading 1:    20pt, #0F4761, space-before 25pt / after 6pt
  Heading 2:    16pt, #0F4761, space-before 11pt / after 6pt
  Heading 3:    14pt, #0F4761
  Page margins: 1 inch all around (1440 twips)
  Page size:    8.5 × 11 in (12240 × 15840 twips)
  Logo:         IMAN black square logo (image1.png from original)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Color palette (exact from original) ───────────────────────────────────────
HEADING_BLUE  = RGBColor(0x0F, 0x47, 0x61)   # headings
BODY_BLACK    = RGBColor(0x0D, 0x0D, 0x0D)   # body text
PLACEHOLDER   = RGBColor(0xB0, 0x60, 0x00)   # amber — easy to spot and find/replace
GUIDANCE_GRAY = RGBColor(0x88, 0x88, 0x88)   # italic guidance notes
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

GARAMOND      = "Garamond"
APTOS_DISPLAY = "Aptos Display"
LOGO_PATH     = "/tmp/iman_logo.png"

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)

# ── Low-level helpers ──────────────────────────────────────────────────────────

def rpr(run, font=GARAMOND, pt=12, bold=False, italic=False, color=BODY_BLACK):
    run.font.name  = font
    run.font.size  = Pt(pt)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color

def ppr(para, before=0, after=8, line=259, indent=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    fmt.line_spacing = Pt(line / 20)   # 259 half-points → 12.95pt
    fmt.alignment    = align
    if indent is not None:
        fmt.left_indent = Inches(indent)

def new_para(doc, before=0, after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    ppr(p, before=before, after=after, align=align)
    return p

# ── Content builders ───────────────────────────────────────────────────────────

def heading1(doc, text):
    """H1: Aptos Display 20pt #0F4761, space-before 25pt."""
    p = doc.add_paragraph()
    ppr(p, before=25, after=6, line=240)
    run = p.add_run(text)
    rpr(run, font=APTOS_DISPLAY, pt=20, color=HEADING_BLUE)
    return p

def heading2(doc, text):
    """H2: Aptos Display 16pt #0F4761."""
    p = doc.add_paragraph()
    ppr(p, before=14, after=6, line=240)
    run = p.add_run(text)
    rpr(run, font=APTOS_DISPLAY, pt=16, color=HEADING_BLUE)
    return p

def body(doc, text, bold=False, color=BODY_BLACK, before=0, after=8):
    p = new_para(doc, before=before, after=after)
    run = p.add_run(text)
    rpr(run, bold=bold, color=color)
    return p

def placeholder(doc, text, before=0, after=8):
    p = new_para(doc, before=before, after=after)
    run = p.add_run(text)
    rpr(run, color=PLACEHOLDER)
    return p

def guidance(doc, text):
    """Italic gray note — delete before submitting."""
    p = new_para(doc, before=0, after=4)
    ppr(p, before=0, after=4)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(f"ⓘ  {text}")
    rpr(run, italic=True, color=GUIDANCE_GRAY, pt=10)
    return p

def divider(doc):
    """Exact match: bold underscores used as section divider in original."""
    p = new_para(doc, before=6, after=6)
    run = p.add_run("_" * 78)
    rpr(run, bold=True, color=BODY_BLACK)
    return p

def focus_area_header(doc, name_placeholder, goal_placeholder):
    """
    Replicates the original's fused title+goal paragraph:
    bold name (large) + italic goal on same paragraph, then divider below.
    Original: name and goal are in the SAME Normal paragraph, name portion bold.
    """
    p = new_para(doc, before=12, after=4)
    r1 = p.add_run(name_placeholder)
    rpr(r1, pt=13, bold=True, color=PLACEHOLDER)
    r2 = p.add_run("  ")
    rpr(r2, pt=12, color=BODY_BLACK)
    r3 = p.add_run(goal_placeholder)
    rpr(r3, pt=12, italic=True, color=PLACEHOLDER)
    return p

def partnerships_label(doc):
    p = new_para(doc, before=8, after=2)
    run = p.add_run("PARTNERSHIPS:")
    rpr(run, bold=True, color=BODY_BLACK)
    return p

def partner_bullet(doc, text, is_placeholder=False):
    p = doc.add_paragraph(style="List Bullet")
    ppr(p, before=0, after=2, line=240)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    rpr(run, pt=12, color=PLACEHOLDER if is_placeholder else BODY_BLACK)
    return p

def budget_subhead(doc, text, is_placeholder=False):
    """ListParagraph 14pt bold — exact match to original budget subheadings."""
    p = doc.add_paragraph(style="List Paragraph")
    ppr(p, before=8, after=4, line=240)
    run = p.add_run(text)
    rpr(run, pt=14, bold=True, color=PLACEHOLDER if is_placeholder else BODY_BLACK)
    return p

def budget_item(doc, text, is_placeholder=False):
    p = doc.add_paragraph(style="List Paragraph")
    ppr(p, before=0, after=4, line=259)
    run = p.add_run(text)
    rpr(run, pt=12, color=PLACEHOLDER if is_placeholder else BODY_BLACK)
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        rpr(run, pt=11, bold=True, color=WHITE)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F4761')
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            rpr(run, pt=11, color=PLACEHOLDER)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# IMAN logo — centred, same size as in original (~1.5 inch square)
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
ppr(p_logo, before=24, after=12)
try:
    run_logo = p_logo.add_run()
    run_logo.add_picture(LOGO_PATH, width=Inches(1.5))
except Exception as e:
    run_logo = p_logo.add_run("[IMAN LOGO]")
    rpr(run_logo, bold=True, color=BODY_BLACK)

# Title block — matches original cover text
p_title = new_para(doc, before=0, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p_title.add_run("Comprehensive Grant Proposal")
rpr(r, font=APTOS_DISPLAY, pt=22, color=HEADING_BLUE)

p_sub = new_para(doc, before=0, after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
r2 = p_sub.add_run("[Funder Name] & IMAN")
rpr(r2, font=GARAMOND, pt=14, color=PLACEHOLDER)

# Cover fields
for label, ph in [
    ("To:       ", "[Funder Name]"),
    ("Subject:  ", "[Program or Initiative Name]"),
    ("Amount:   ", "$[X]M over [X] years"),
    ("Date:     ", "[Month, Year]"),
    ("Contact:  ", "[IMAN Contact Name, Title, email]"),
]:
    p = new_para(doc, before=0, after=3)
    r1 = p.add_run(label)
    rpr(r1, bold=True, color=BODY_BLACK)
    r2 = p.add_run(ph)
    rpr(r2, color=PLACEHOLDER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual — matches original TOC structure)
# ══════════════════════════════════════════════════════════════════════════════

p_toc_h = new_para(doc, before=0, after=8)
r = p_toc_h.add_run("Contents")
rpr(r, font=APTOS_DISPLAY, pt=16, bold=True, color=HEADING_BLUE)

toc = [
    ("SUMMARY", False),
    ("BACKGROUND & CONTEXT", False),
    ("~ IMAN 2027 ~", False),
    ("    [Focus Area 1]", True),
    ("    [Focus Area 2]", True),
    ("    [Focus Area 3]", True),
    ("    [Focus Area 4]", True),
    ("    [Focus Area 5]", True),
    ("BUDGET & SUSTAINABILITY", False),
    ("[NEIGHBORHOOD]: LEADERSHIP & COALITION", True),
    ("SPIRITUAL & CULTURAL IMPERATIVE OF THE WORK", False),
    ("CASE STUDIES & VALIDATION FOR THE WORK", False),
]
for text, is_ph in toc:
    p = new_para(doc, before=0, after=2)
    run = p.add_run(text)
    rpr(run, pt=12, color=PLACEHOLDER if is_ph else BODY_BLACK)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1 — SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "SUMMARY")
guidance(doc, "2–3 paragraphs. Open with IMAN's tenure and the core disparity statistic. Name the ask — dollar amount, duration, funder — in paragraph 3. Close with the long-term vision.")

placeholder(doc,
    "For over [X] years, the Inner-City Muslim Action Network (IMAN) has been a driving force in "
    "addressing systemic disparities through [health / housing / workforce / food justice] initiatives "
    "in underserved urban areas. [Cite the key disparity statistic — e.g., a 2019 study revealed a "
    "staggering 30-year life expectancy gap between Englewood and the downtown Streeterville neighborhood, "
    "highlighting the urgency of IMAN's mission.]"
)
placeholder(doc,
    "Through its [Year] strategic goals, IMAN and its partners aim to [core mission statement] by "
    "addressing [list 3–5 domains: healthcare, housing, workforce development, food justice, "
    "neighborhood revitalization], while developing a long-term vision to systematically and "
    "quantitatively reduce disparities."
)
placeholder(doc,
    "To support this effort, IMAN seeks a [X]-year, $[X] million [grant type] from [Funder Name]. "
    "This grant will enable us to [meet specific strategic objectives] and lay the foundation for a "
    "bold [X]-year initiative to [long-term vision — e.g., transform life expectancy and quality of life "
    "for thousands of residents, advancing equity and creating lasting opportunity for the communities "
    "most impacted by historic disinvestment]."
)

# ══════════════════════════════════════════════════════════════════════════════
#  2 — BACKGROUND & CONTEXT
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "BACKGROUND & CONTEXT:")
guidance(doc, "2–3 paragraphs. Describe IMAN's integrated approach. Show the trajectory: small credibility → large-scale transformation. Include specific dollar figures for past investments.")

placeholder(doc,
    "IMAN has developed a comprehensive, integrative approach to address deeply ingrained structural "
    "and systemic inequalities. Targeting [primary populations — e.g., justice-involved, immigrant, "
    "and other vulnerable communities], IMAN's strategy also emphasizes [secondary elements — e.g., "
    "arts programming, community organizing, and advocacy] at the heart of its initiatives."
)
placeholder(doc,
    "[Describe 2–3 landmark past investments with dollar amounts and outcomes. E.g.: 'In [year], as "
    "part of [campaign name], IMAN [action]. Though [challenge], IMAN and its partners persevered, "
    "turning what was initially a $[X] award into a $[X]M investment. This includes [list key assets].']"
)
placeholder(doc,
    "[Describe the current strategic plan name, timeframe, and how this grant fits into it. E.g.: "
    "'Building on the momentum of [X] decades of work in [neighborhoods], IMAN's [Year] strategic "
    "plan aims to deepen its impact through five ambitious focus areas. The comprehensive grant "
    "proposal to [Funder] is intended to support these growth areas.']"
)

# Strategic plan header — original uses a Heading1-style centered section break
p_plan = new_para(doc, before=16, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p_plan.add_run("~ [STRATEGIC PLAN NAME] ~")
rpr(r, font=APTOS_DISPLAY, pt=20, color=HEADING_BLUE)

body(doc,
    "[1–2 sentences introducing the five focus areas and their relationship to the strategic plan "
    "and this grant request.]",
    color=PLACEHOLDER
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  3 — FOCUS AREAS  (5 × identical block structure)
# ══════════════════════════════════════════════════════════════════════════════

focus_areas = [
    ("[Focus Area 1 Name — e.g., IMAN's Holistic Health Center]",
     "[Goal statement: one sentence. Action verb + what will be built + for whom + to achieve what. "
      "E.g.: Build an integrated FQHC facility at 63rd/California to expand healthcare delivery and "
      "improve patient outcomes for underserved Chicagoland communities.]"),

    ("[Focus Area 2 Name — e.g., Holistic ReEntry]",
     "[Goal statement. E.g.: Drive neighborhood stabilization & expand IMAN's holistic reentry "
      "apparatus by transforming [building] into a permanent supportive housing hub, while strengthening "
      "workforce development, transitional housing, and wrap-around support.]"),

    ("[Focus Area 3 Name — e.g., Food Ecosystems]",
     "[Goal statement. E.g.: Drive innovative solutions for food justice while expanding community "
      "access to nutritious foods and nutrition initiatives.]"),

    ("[Focus Area 4 Name — e.g., Vibrant Community Nodes]",
     "[Goal statement. E.g.: Foster vibrancy, safety, and cultural vitality at [intersection 1] "
      "and [intersection 2] nodes.]"),

    ("[Focus Area 5 Name — e.g., Model Building]",
     "[Goal statement. E.g.: Forge durable strategic partnerships around a shared vision and "
      "formalize IMAN's impact model for scalable and sustainable success.]"),
]

for i, (name, goal) in enumerate(focus_areas):
    focus_area_header(doc, name, goal)
    guidance(doc,
        "Lead with the most impressive outcome metrics. Use unduplicated counts. "
        "Include year-over-year growth wherever possible."
    )
    placeholder(doc,
        "[Program name] [provides / operates / supports] [brief description of the program]. "
        "In [year], we served [#] unduplicated [patients / participants / residents] through "
        "[#] [visits / sessions / interactions]. Key outcomes:"
    )
    for metric in [
        "[Metric 1 with number and year-over-year comparison — e.g., '91% retention rate in 2024']",
        "[Metric 2 — e.g., '477 HIV screenings, a 48% increase over 2023']",
        "[Metric 3 — e.g., '$1.7M in sales, keeping wages in the community']",
    ]:
        partner_bullet(doc, metric, is_placeholder=True)

    # Expansion sub-section
    p_exp = new_para(doc, before=10, after=4)
    r1 = p_exp.add_run("Moving forward, IMAN will expand [program] in [X] key ways:")
    rpr(r1, color=PLACEHOLDER)

    # Sub-initiative 1
    p_s1 = new_para(doc, before=4, after=4)
    r_label = p_s1.add_run("[Sub-initiative 1 Name]: ")
    rpr(r_label, bold=True, color=BODY_BLACK)
    r_desc = p_s1.add_run(
        "[Description of what will be built, converted, or launched. Include cost ($[X]M), "
        "scale (sq ft / units / jobs), design partners, and any policy recognition or awards.]"
    )
    rpr(r_desc, color=PLACEHOLDER)

    p_s2 = new_para(doc, before=4, after=2)
    r_label2 = p_s2.add_run("[Sub-initiative 2 Name]: ")
    rpr(r_label2, bold=True, color=BODY_BLACK)
    r_desc2 = p_s2.add_run("[Description.] Specifically, this includes:")
    rpr(r_desc2, color=PLACEHOLDER)

    for ws in [
        "[Workstream 1 — e.g., Workforce Development: doubling capacity to 100 individuals/year, adding 3 new trades]",
        "[Workstream 2 — e.g., Housing Expansion: increasing transitional housing from 20 to 52 units]",
        "[Workstream 3 — e.g., Wraparound Services: deepening case management, legal aid, food access]",
    ]:
        p_ws = new_para(doc, before=0, after=2)
        p_ws.paragraph_format.left_indent = Inches(0.3)
        r = p_ws.add_run(ws)
        rpr(r, bold=True, color=PLACEHOLDER)

    partnerships_label(doc)
    guidance(doc,
        "List 5–10 named partners. Include government agencies, hospitals, foundations, and community "
        "organizations. Named partners reduce funder risk perception."
    )
    for j in range(1, 7):
        partner_bullet(doc, f"[Partner {j} — full organization name]", is_placeholder=True)

    if i < len(focus_areas) - 1:
        divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  4 — BUDGET & SUSTAINABILITY
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "BUDGET & SUSTAINABILITY")
guidance(doc,
    "Narrate the path to sustainability across 3 categories. Answer: 'What happens after the grant "
    "ends?' The original had no budget table in the narrative — it referenced an attached budget doc."
)

placeholder(doc,
    "Funding from [Funder Name] will provide a stable platform from which IMAN will launch and expand "
    "these transformative community initiatives while establishing a path to financial sustainability. "
    "IMAN is pursuing three categories of revenue to ensure the continued impact and growth of these "
    "assets beyond [grant end year]:"
)

body(doc, "Earned Revenue", bold=True, before=8, after=4)
guidance(doc, "14pt bold in original. Describe each earned revenue stream with trajectory and projections.")

for item in [
    ("[Revenue Stream 1 — e.g., Health Center revenue]: "
     "[Trajectory. E.g.: 'The expanded Holistic Health Center is projected to increase patient "
     "revenues by more than ninefold, reaching $9M by 2030 while creating 30 new full-time jobs. "
     "In the interim, IMAN will also work with a financial modeling consultant to increase "
     "reimbursement for visits and marketing to increase patient encounters.']"),

    ("[Revenue Stream 2 — e.g., Housing income]: "
     "[When it activates. E.g.: 'In [year] and [year+1], IMAN will see a significant increase in "
     "revenue from housing vouchers and rental income, with the opening of [X] units at [project name] "
     "and the buildout of new transitional housing.']"),

    ("[Revenue Stream 3 — e.g., Retail / Facility Rental]: "
     "[Growth trajectory. E.g.: 'Sales at the Fresh Market continue to grow by double digits year "
     "over year through word of mouth alone. In [year], IMAN will invest in marketing and paid "
     "advertising to continue this trajectory while planning a store expansion in the next five years.']"),

    ("[Revenue Stream 4 — e.g., Facility rental / café tenant]: "
     "[E.g.: 'The [space name] will feature a café tenant and offer attractive meeting spaces for "
     "community organizations and partners at competitive rates. Additional facilities currently "
     "provided at no cost will transition into rental revenue sources as needed.']"),
]:
    budget_item(doc, item, is_placeholder=True)

budget_subhead(doc,
    "Grant Funding — [Funder Name] will continue working with our sustaining partners and donors "
    "while pursuing new grant funding sources, including [list 2–3 target funders or fund types].",
    is_placeholder=True
)

budget_subhead(doc,
    "Operational Optimization: As IMAN's operations continue to grow and evolve, we will evaluate, "
    "structure, and optimize our operations to manage high-impact programming efficiently and effectively.",
    is_placeholder=False
)

placeholder(doc,
    "The accompanying budget document outlines the capital projects within each focus area, detailing "
    "their current sources, uses, and funding gaps. In a separate tab titled 'Comprehensive Grant "
    "Budget,' we have listed the programmatic and capacity-building support required to achieve these "
    "goals. This tab provides a detailed breakdown of the resources necessary to implement these "
    "initiatives and ensure their effective and successful execution."
)

# ══════════════════════════════════════════════════════════════════════════════
#  5 — COMMUNITY CONTEXT
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "[NEIGHBORHOOD]: LEADERSHIP & COALITION")
guidance(doc,
    "Frame the neighborhood as resilient and in momentum — not deficient. Use CMAP or comparable data. "
    "Name cultural leaders and list catalytic investments already underway. Show IMAN is not alone."
)

placeholder(doc,
    "[Neighborhood name] is a community rich in history, resilience, and untapped potential. While it "
    "has faced challenges — including [name structural challenges, e.g., complexities of multiple "
    "aldermanic wards and the impact of systemic disinvestment] — community organizations and leaders "
    "are building a unified vision for [neighborhood]'s revitalization."
)
placeholder(doc,
    "The data showcases the strength of the community. Per [data source — e.g., CMAP's 2024 community "
    "data snapshot], [neighborhood] has shown significant progress despite facing challenges. "
    "[Metric 1 — e.g., adjusted for inflation, the median household income rose from $[X] to $[X].] "
    "[Metric 2 — e.g., unemployment fell from [X]% to [X]% over the last decade.] "
    "[Metric 3 — e.g., the percentage of residents holding bachelor's degrees increased from [X]% to [X]%.]"
)
placeholder(doc,
    "[Name 2–3 cultural leaders or organizations raising the neighborhood's visibility and "
    "fueling a virtuous cycle of public and private investment. E.g.: artists, historians, "
    "neighborhood collectives.]"
)
placeholder(doc,
    "[List 2–4 major catalytic investments already underway with dollar figures. E.g.: 'Significant "
    "initiatives, including the $40M transformation of [school/building] into affordable housing, "
    "the $33M conversion of [building] into permanent supportive housing, and the $40M [project name] "
    "mixed-income housing project, herald a fresh chapter of growth and potential in the neighborhood.']"
)
placeholder(doc,
    "[1–2 sentences on community governance and how IMAN navigates it. E.g.: 'By engaging both "
    "aldermen as key partners, the community aims to create a cohesive vision that reflects the "
    "entire community's aspirations.'] With your support, we can accelerate this progress and ensure "
    "[neighborhood] thrives as a vibrant, unified community."
)

# ══════════════════════════════════════════════════════════════════════════════
#  6 — SPIRITUAL & CULTURAL IMPERATIVE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "SPIRITUAL & CULTURAL IMPERATIVE OF THE WORK")
guidance(doc,
    "IMAN's most distinctive section. Don't make it generic. Connect Muslim values to the multifaith "
    "coalition. Name the current political/cultural moment of urgency directly — the original named "
    "the Muslim-Jewish partnership and the 2024 elections explicitly. This differentiates IMAN from "
    "every other nonprofit proposal."
)

placeholder(doc,
    "Since its inception, IMAN has carved a distinct niche within the wider social justice movement "
    "in urban areas. It is deeply anchored in its Muslim values while also drawing inspiration from "
    "various faith traditions, sacred practices, and cultural expressions that resonate with its "
    "mission of community transformation. Through that framework, the organization has served as a "
    "convener, national model, and thought leader in the larger field."
)
placeholder(doc,
    "Through community organizing campaigns that contend with longstanding issues of inequity and "
    "injustice, IMAN centers its work around communal healing as a spiritual practice. [Describe a "
    "specific example of IMAN's multifaith coalition work — name the partners, the outcome, and its "
    "lasting significance. E.g.: the MLK memorial designed and built at IMAN, bringing together "
    "Muslim, Jewish, and Christian communities alongside organizations from across the city.]"
)
placeholder(doc,
    "[Connect the current political, cultural, or communal moment to the urgency of this specific "
    "partnership. Be direct. Name the context. E.g.: 'The recent [elections / events] underscore the "
    "uncertainty, vulnerability, and concerns of communities dependent on essential services, support, "
    "and care. This moment reinforces the spiritual necessity and urgency of the partnerships that "
    "[funder] and IMAN represent.'] "
    "IMAN continues to see this work as profoundly significant spiritual work — a model that has been "
    "discussed, researched, and visited by people worldwide and that we believe has the power to "
    "demonstrate how people and institutions with important shared values can have a profound impact "
    "on seemingly intractable issues, even amid intensely challenging circumstances."
)

# ══════════════════════════════════════════════════════════════════════════════
#  7 — CASE STUDIES & VALIDATION
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "CASE STUDIES & VALIDATION FOR THE WORK:")
guidance(doc,
    "Open with an academic/expert quote from a named university partner. List 3–5 comparable models — "
    "each validating a different IMAN focus area. The original used a Northwestern Medicine quote. "
    "Attach full case studies as an appendix."
)

# Expert quote block
p_q = new_para(doc, before=8, after=8)
p_q.paragraph_format.left_indent = Inches(0.4)
r_q = p_q.add_run(
    "\"[Academic quote validating the multilevel / comprehensive approach — get a direct quote from "
    "a named university partner (Northwestern, UChicago, Rush, Cornell, etc.). E.g.: 'Multilevel "
    "strategies can more effectively and sustainably address complex health issues, such as health "
    "disparities, by approaching them from various angles within environmental systems. This method "
    "targets both the root causes and immediate health requirements.' — [Dr. Name, Title, Institution]]\""
)
rpr(r_q, italic=True, color=PLACEHOLDER)

placeholder(doc,
    "Other examples of high-impact models taking a multilevel approach locally and nationally include:"
)

add_table(doc,
    ["Organization", "Location", "Focus Area", "Relevance to IMAN's Model"],
    [
        ["[Org 1]", "[City, State]",
         "[e.g., ReEntry + workforce + housing]",
         "[How it validates IMAN's comprehensive approach]"],
        ["[Org 2]", "[City, State]",
         "[e.g., Cradle-to-career youth pathways]",
         "[How it validates IMAN's model]"],
        ["[Org 3]", "[City, State]",
         "[e.g., HUD Choice Neighborhoods]",
         "[How it validates IMAN's model]"],
        ["[Org 4]", "[City, State]",
         "[e.g., Nonprofit grocery / food retail]",
         "[How it validates IMAN's Fresh Market model]"],
    ]
)

placeholder(doc,
    "Enclosed with this report are case studies of these and other organizations making deep and "
    "lasting change in communities through a comprehensive approach. Also included are summaries of "
    "relevant academic journal articles and intersections with IMAN's work."
)

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDICES CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "APPENDICES (ATTACH SEPARATELY)")

for item in [
    "Capital budget — sources, uses, and funding gaps per focus area",
    "Comprehensive Grant Budget — programmatic and capacity-building costs by initiative",
    "Case studies (3–5 comparable organizations)",
    "Academic article summaries relevant to each focus area",
    "Audited financials (most recent 2 years)",
    "IRS 501(c)(3) determination letter",
    "Board list",
    "Key staff bios",
]:
    partner_bullet(doc, f"☐  {item}")

# ══════════════════════════════════════════════════════════════════════════════
#  WRITING GUIDANCE (final page)
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading1(doc, "WRITING GUIDANCE")
guidance(doc, "Delete this entire page before submitting. For internal use only.")

tips = [
    ("On metrics:",
     "Always pair a number with a comparison — year-over-year growth, vs. city/state average, or vs. "
     "IMAN's own prior baseline. Standalone numbers don't land. Unduplicated counts only."),
    ("On partnerships:",
     "'Northwestern Memorial Hospital' is more persuasive than 'a major academic medical center.' "
     "Named partners signal credibility and reduce perceived funder risk. List 5–10 per focus area."),
    ("On the ask:",
     "State the dollar amount and duration in paragraph 3 of the Summary. Don't bury it. The original "
     "Crown-IMAN ask: two-year, $5M, named in the first page."),
    ("On length:",
     "The Crown-IMAN proposal ran ~15 pages of narrative. Cut anything that doesn't: (1) establish "
     "credibility, (2) quantify impact, or (3) reduce perceived risk."),
    ("On the focus area header:",
     "In the original, the focus area name and goal statement appear in the SAME paragraph — name "
     "bold/large, then goal italic immediately after. Keep this fused format."),
    ("On the Spiritual & Cultural section:",
     "Don't make it generic. The Crown-IMAN proposal named the Palestinian-Jewish dynamic and the "
     "2024 elections directly. Name the current equivalent moment. This is the section that "
     "differentiates IMAN from every other nonprofit proposal."),
    ("On tone:",
     "IMAN's voice is simultaneously urgent and grounded. Avoid deficit framing ('the community is "
     "suffering') and triumphalism ('IMAN is transforming everything'). Target tone: 'We have earned "
     "the right to make this ask, and here is exactly how we will use it.'"),
    ("Amber text = placeholder fields.",
     "Use Find & Replace (Ctrl+H) to locate all amber/orange text. Every field must be filled before "
     "submission. Delete all ⓘ guidance notes."),
]

for label, text in tips:
    p = new_para(doc, before=6, after=4)
    r1 = p.add_run(f"{label}  ")
    rpr(r1, bold=True, color=HEADING_BLUE)
    r2 = p.add_run(text)
    rpr(r2, color=BODY_BLACK)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/workspaces/IMAN/Data library/Consulting/Week 3/Grant Proposal Template.docx"
doc.save(out)
print(f"✓ Saved: {out}")
